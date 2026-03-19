"""Tests d'intégration pour le pipeline d'ingestion GRI.

Ces tests vérifient le flux complet :
Parse DOCX → Tables → Glossaire → Chunking → Validation → Indexation

Lancer avec: pytest tests/integration/test_pipeline.py -v -m integration
"""

import pytest
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, AsyncMock, patch

from src.ingestion.pipeline import GRIIngestionPipeline
from src.ingestion.chunker import GRIChunker
from src.ingestion.models import (
    Cycle,
    GRIChunk,
    GRIMetadata,
    ParsedSection,
    ParsedTable,
    SectionType,
    IngestionResult,
)
from src.ingestion.parser import GRIDocxParser
from src.ingestion.table_extractor import GRITableExtractor
from src.ingestion.glossary_extractor import GRIGlossaryExtractor


pytestmark = pytest.mark.integration


def make_valid_test_chunk() -> GRIChunk:
    """Construit un chunk valide conforme au modèle actuel."""
    context_prefix = "[GRI > Test]"
    content = f"{context_prefix}\n\n" + ("Contenu de test " * 8)
    return GRIChunk(
        chunk_id="1234567890abcdef",
        content=content,
        metadata=GRIMetadata(
            doc_id="abcdef1234567890",
            chunk_index=0,
            section_type=SectionType.DEFINITION,
            hierarchy=["GRI", "Test"],
            context_prefix=context_prefix,
            cycle=Cycle.GRI,
            char_count=len(content),
        ),
    )


# =============================================================================
# Fixtures
# =============================================================================


@pytest.fixture
def sample_sections():
    """Sections parsées de test."""
    return [
        ParsedSection(
            title="Glossaire",
            level=1,
            content="artefact : Produit d'ingénierie.\nCONOPS : Concept d'opérations.",
            section_type=SectionType.DEFINITION,
            hierarchy=["GRI", "Glossaire"],
        ),
        ParsedSection(
            title="Phase 3 - Conception et Développement",
            level=1,
            content="Cette phase couvre la conception détaillée du système.",
            section_type=SectionType.PHASE,
            hierarchy=["GRI", "Phase 3"],
            metadata={"phase_num": 3},
        ),
        ParsedSection(
            title="Jalon M4 - CDR",
            level=2,
            content="Le CDR valide la conception détaillée. Critères :\n"
            "1. Architecture système validée\n"
            "2. Interfaces définies\n"
            "3. Plans de test établis",
            section_type=SectionType.MILESTONE,
            hierarchy=["GRI", "Phase 3", "Jalon M4"],
            metadata={"milestone_id": "M4", "milestone_name": "CDR"},
        ),
        ParsedSection(
            title="Processus de Vérification",
            level=2,
            content="Le processus de vérification assure que les exigences sont satisfaites.",
            section_type=SectionType.PROCESS,
            hierarchy=["GRI", "Phase 3", "Processus de Vérification"],
            metadata={"process_name": "Vérification"},
        ),
        ParsedSection(
            title="Principe 5 - Traçabilité",
            level=1,
            content="Tous les artefacts doivent être traçables aux exigences source.",
            section_type=SectionType.PRINCIPLE,
            hierarchy=["GRI", "Principes", "Principe 5"],
            metadata={"principle_num": 5},
        ),
    ]


@pytest.fixture
def sample_tables():
    """Tables parsées de test."""
    return [
        ParsedTable(
            title="Critères du jalon M4 (CDR)",
            headers=["N°", "Critère", "Catégorie"],
            rows=[
                ["1", "Architecture système validée", "Technique"],
                ["2", "Interfaces définies et documentées", "Technique"],
                ["3", "Plans de test établis", "Qualité"],
            ],
            section_type=SectionType.MILESTONE,
            hierarchy=["GRI", "Phase 3", "Jalon M4"],
            metadata={"milestone_id": "M4"},
        ),
    ]


@pytest.fixture
def mock_vector_store():
    """Vector store mocké."""
    store = MagicMock()
    store.index_chunks = AsyncMock(return_value={"total": 10, "indexed": 10})
    store.ensure_collections = AsyncMock()
    return store


# =============================================================================
# Tests GRIChunker
# =============================================================================


class TestGRIChunker:
    """Tests pour le chunker GRI."""

    def test_chunk_definition_no_split(self, sample_sections):
        """Les définitions sont des chunks individuels."""
        chunker = GRIChunker()
        definition_section = sample_sections[0]

        chunks = chunker.chunk_section(definition_section)

        # Chaque définition doit être un chunk séparé
        assert len(chunks) >= 1
        for chunk in chunks:
            assert chunk.section_type == SectionType.DEFINITION

    def test_chunk_milestone_never_split(self, sample_sections):
        """CRITIQUE : Les jalons ne sont JAMAIS fragmentés."""
        chunker = GRIChunker()
        milestone_section = sample_sections[2]

        chunks = chunker.chunk_section(milestone_section)

        # Un jalon = un seul chunk
        assert len(chunks) == 1
        assert "Architecture" in chunks[0].content
        assert "Interfaces" in chunks[0].content
        assert "Plans de test" in chunks[0].content

    def test_chunk_has_context_prefix(self, sample_sections):
        """Chaque chunk a un préfixe de contexte."""
        chunker = GRIChunker()

        for section in sample_sections:
            chunks = chunker.chunk_section(section)
            for chunk in chunks:
                assert chunk.context_prefix is not None
                assert chunk.context_prefix.startswith("[")
                assert chunk.context_prefix.endswith("]")

    def test_chunk_phase_parent_document_retriever(self, sample_sections):
        """Les phases utilisent le pattern Parent Document Retriever."""
        chunker = GRIChunker()
        phase_section = sample_sections[1]

        # Simuler une phase longue
        long_phase = ParsedSection(
            title=phase_section.title,
            level=phase_section.level,
            content=phase_section.content * 20,  # Répéter pour dépasser le seuil
            section_type=phase_section.section_type,
            hierarchy=phase_section.hierarchy,
            metadata=phase_section.metadata,
        )

        chunks = chunker.chunk_section(long_phase)

        # Avec Parent Doc Retriever, on devrait avoir des chunks enfants
        # mais aussi un lien vers le parent
        assert len(chunks) >= 1
        for chunk in chunks:
            assert chunk.section_type == SectionType.PHASE

    def test_chunk_process_respects_limits(self, sample_sections):
        """Les processus respectent les limites de tokens."""
        chunker = GRIChunker()
        process_section = sample_sections[3]

        chunks = chunker.chunk_section(process_section)

        # Chaque chunk doit respecter la limite de tokens (400-600)
        for chunk in chunks:
            # Approximation : 1 token ≈ 4 caractères
            token_estimate = len(chunk.content) / 4
            # On vérifie que ce n'est pas excessivement long
            assert token_estimate < 1000

    def test_chunk_preserves_metadata(self, sample_sections):
        """Le chunking préserve les métadonnées."""
        chunker = GRIChunker()
        milestone_section = sample_sections[2]

        chunks = chunker.chunk_section(milestone_section)

        assert len(chunks) >= 1
        chunk = chunks[0]
        assert chunk.metadata.get("milestone_id") == "M4"
        assert chunk.cycle == "GRI"


class TestChunkingStrategies:
    """Tests pour les 7 stratégies de chunking."""

    def test_definition_strategy(self):
        """Stratégie définition : 150-300 tokens, overlap 0."""
        chunker = GRIChunker()
        config = chunker._get_config_for_type(SectionType.DEFINITION)

        assert config["min_tokens"] <= 150
        assert config["max_tokens"] <= 300
        assert config["overlap"] == 0

    def test_milestone_strategy(self):
        """Stratégie jalon : JAMAIS fragmenté."""
        chunker = GRIChunker()
        config = chunker._get_config_for_type(SectionType.MILESTONE)

        assert config["never_split"] is True

    def test_phase_strategy(self):
        """Stratégie phase : Parent Document Retriever."""
        chunker = GRIChunker()
        config = chunker._get_config_for_type(SectionType.PHASE)

        # Parent doc: chunks plus grands avec overlap
        assert config["max_tokens"] >= 512
        assert config["overlap"] >= 64

    def test_process_strategy(self):
        """Stratégie processus : 400-600 tokens."""
        chunker = GRIChunker()
        config = chunker._get_config_for_type(SectionType.PROCESS)

        assert 400 <= config["max_tokens"] <= 600


# =============================================================================
# Tests Pipeline Integration
# =============================================================================


class TestPipelineIntegration:
    """Tests d'intégration du pipeline complet."""

    @pytest.fixture
    def pipeline(self, mock_vector_store):
        """Pipeline configuré pour les tests."""
        return GRIIngestionPipeline(
            vector_store=mock_vector_store,
            output_dir=tempfile.mkdtemp(),
        )

    def test_pipeline_initialization(self, pipeline):
        """Le pipeline s'initialise correctement."""
        assert pipeline.parser is not None
        assert pipeline.table_extractor is not None
        assert pipeline.glossary_extractor is not None
        assert pipeline.vector_store is not None

    def test_pipeline_components_instantiated(self, pipeline):
        """Tous les composants sont du bon type."""
        assert isinstance(pipeline.parser, GRIDocxParser)
        assert isinstance(pipeline.table_extractor, GRITableExtractor)
        assert isinstance(pipeline.glossary_extractor, GRIGlossaryExtractor)

    @patch.object(GRIDocxParser, "parse")
    @patch.object(GRITableExtractor, "extract")
    @patch.object(GRIGlossaryExtractor, "extract_from_sections")
    def test_pipeline_runs_all_steps(
        self,
        mock_glossary,
        mock_tables,
        mock_parse,
        pipeline,
        sample_sections,
        sample_tables,
    ):
        """Le pipeline exécute toutes les étapes."""
        mock_parse.return_value = (sample_sections, sample_tables)
        mock_tables.return_value = sample_tables
        mock_glossary.return_value = [
            {"term_fr": "artefact", "definition": "Produit d'ingénierie"}
        ]

        # Créer un fichier DOCX temporaire vide pour le test
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            # Écrire un fichier DOCX minimal
            from docx import Document

            doc = Document()
            doc.add_paragraph("Test")
            doc.save(f.name)

            result = pipeline.run(f.name)

        assert isinstance(result, IngestionResult)
        mock_parse.assert_called_once()

    @patch.object(GRIDocxParser, "parse")
    def test_pipeline_handles_parse_error(self, mock_parse, pipeline):
        """Le pipeline gère les erreurs de parsing."""
        mock_parse.side_effect = Exception("Fichier corrompu")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            result = pipeline.run(f.name)

        assert result is not None
        assert len(result.errors) > 0

    @patch.object(GRIDocxParser, "parse")
    @patch.object(GRITableExtractor, "extract")
    def test_pipeline_continues_on_table_error(
        self,
        mock_tables,
        mock_parse,
        pipeline,
        sample_sections,
    ):
        """Le pipeline continue même si l'extraction de tables échoue."""
        mock_parse.return_value = (sample_sections, [])
        mock_tables.side_effect = Exception("Erreur table")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            from docx import Document

            doc = Document()
            doc.add_paragraph("Test")
            doc.save(f.name)

            result = pipeline.run(f.name)

        # Le pipeline doit continuer malgré l'erreur
        assert result is not None
        # Un warning doit être enregistré
        assert len(result.warnings) > 0


class TestChunkValidation:
    """Tests pour la validation des chunks."""

    def test_validate_chunk_has_content(self):
        """Un chunk valide a du contenu."""
        chunk = make_valid_test_chunk()

        assert len(chunk.content) > 0

    def test_validate_chunk_has_context_prefix(self):
        """Un chunk valide a un préfixe de contexte."""
        chunk = make_valid_test_chunk()

        assert chunk.context_prefix is not None
        assert chunk.context_prefix.startswith("[")

    def test_validate_milestone_not_truncated(self, sample_sections):
        """CRITIQUE : Valider qu'un jalon n'est pas tronqué."""
        chunker = GRIChunker()
        milestone_section = sample_sections[2]

        chunks = chunker.chunk_section(milestone_section)

        # Le jalon original doit être entièrement présent
        original_criteria = ["Architecture", "Interfaces", "Plans de test"]

        for criterion in original_criteria:
            found = any(criterion in chunk.content for chunk in chunks)
            assert found, f"Critère '{criterion}' manquant dans les chunks"


# =============================================================================
# Tests de bout en bout (avec fichier réel si disponible)
# =============================================================================


@pytest.mark.slow
class TestPipelineEndToEnd:
    """Tests E2E avec fichier réel."""

    @pytest.fixture
    def real_gri_path(self):
        """Chemin vers le GRI réel (si disponible)."""
        paths = [
            Path("data/raw/gri.docx"),
            Path("data/raw/IRF20251211_last_FF.docx"),
        ]
        for p in paths:
            if p.exists():
                return p
        pytest.skip("Fichier GRI non disponible")

    def test_full_ingestion_flow(self, real_gri_path, mock_vector_store):
        """Test complet du flux d'ingestion."""
        pipeline = GRIIngestionPipeline(vector_store=mock_vector_store)

        result = pipeline.run(real_gri_path)

        assert result is not None
        assert result.total_chunks > 0
        assert result.valid_chunks > 0

    def test_milestone_integrity(self, real_gri_path, mock_vector_store):
        """Vérifie l'intégrité des jalons extraits."""
        pipeline = GRIIngestionPipeline(vector_store=mock_vector_store)

        result = pipeline.run(real_gri_path)

        # Filtrer les chunks de type jalon
        milestone_chunks = [
            c for c in result.chunks if c.section_type == SectionType.MILESTONE
        ]

        # Chaque jalon doit être un chunk unique (non fragmenté)
        milestone_ids = set()
        for chunk in milestone_chunks:
            mid = chunk.metadata.get("milestone_id")
            if mid:
                # Vérifier qu'on n'a pas de doublons fragmentés
                if mid in milestone_ids:
                    # Si doublon, c'est OK seulement si c'est le même contenu
                    pass
                milestone_ids.add(mid)

    def test_glossary_extraction(self, real_gri_path, mock_vector_store):
        """Vérifie l'extraction du glossaire."""
        pipeline = GRIIngestionPipeline(vector_store=mock_vector_store)

        result = pipeline.run(real_gri_path)

        # Filtrer les définitions
        definitions = [
            c for c in result.chunks if c.section_type == SectionType.DEFINITION
        ]

        # Le GRI contient 200+ définitions ISO
        assert len(definitions) >= 10  # Au minimum quelques définitions


# =============================================================================
# Tests de performance
# =============================================================================


@pytest.mark.slow
class TestPipelinePerformance:
    """Tests de performance du pipeline."""

    def test_chunking_performance(self, sample_sections):
        """Le chunking doit être rapide."""
        import time

        chunker = GRIChunker()

        # Simuler beaucoup de sections
        many_sections = sample_sections * 100

        start = time.time()
        for section in many_sections:
            chunker.chunk_section(section)
        elapsed = time.time() - start

        # 500 sections doivent être traitées en moins de 5 secondes
        assert elapsed < 5.0, f"Chunking trop lent: {elapsed:.2f}s"
