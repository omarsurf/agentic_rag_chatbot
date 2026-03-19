"""Parser DOCX pour le document GRI.

Extrait la hiérarchie complète (7 niveaux) et sépare les flux :
- Texte courant (sections)
- Tables
- Glossaire
"""

import re
from pathlib import Path

from docx import Document as DocxDocument
from docx.document import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.core.logging import get_logger
from src.core.milestone_utils import extract_milestone_id
from src.ingestion.models import ParsedSection, ParsedTable, SectionType

log = get_logger(__name__)


class GRIDocxParser:
    """Parser spécialisé pour le document GRI/FAR.

    Préserve la hiérarchie à 7 niveaux et détecte automatiquement
    les types de sections (définition, principe, phase, jalon, etc.).
    """

    # Patterns pour détecter le type de section
    SECTION_TYPE_PATTERNS: dict[str, re.Pattern] = {
        "definition": re.compile(
            r"(Liste\s+des\s+abréviations|Terminologie|Définitions?|Glossaire|"
            r"Termes\s+et\s+définitions|Abbreviations?|Acronymes?|Lexique|"
            r"Index\s+terminologique|Table\s+des\s+(?:abréviations|termes)|"
            r"Sigles\s+et\s+acronymes|Terms\s+and\s+definitions)",
            re.IGNORECASE,
        ),
        "principle": re.compile(
            r"Principe\s+N°\s*\d+",
            re.IGNORECASE,
        ),
        "phase": re.compile(
            r"Phase\s+[1-7]|Phase\s+d['']",
            re.IGNORECASE,
        ),
        "milestone": re.compile(
            r"(Jalon\s+[MJ]\d|Critères\s+(du|de)\s+passage|"
            r"\b(ASR|MNS|SRR|SFR|PDR|CDR|IRR|TRR|PRR|ORR|MNR|SAR)\b)",
            re.IGNORECASE,
        ),
        "process": re.compile(
            r"(Processus|Process)\s+.*(IS\s*15288|Version)",
            re.IGNORECASE,
        ),
        "cir": re.compile(
            r"(CIR|Cycle\s+d['']Innovation\s+Rapide|Phase\s+[1-4].*J\d)",
            re.IGNORECASE,
        ),
    }

    # Patterns pour extraire les identifiants
    PHASE_NUM_PATTERN = re.compile(r"Phase\s+(\d+)", re.IGNORECASE)
    PRINCIPLE_NUM_PATTERN = re.compile(r"Principe\s+N°\s*(\d+)", re.IGNORECASE)

    def __init__(self) -> None:
        self._current_hierarchy: list[str] = ["GRI"]
        self._current_levels: dict[int, str] = {}

    def parse(self, docx_path: str | Path) -> tuple[list[ParsedSection], list[ParsedTable]]:
        """Parse le document DOCX et retourne sections + tables.

        Args:
            docx_path: Chemin vers le fichier DOCX

        Returns:
            Tuple (sections, tables)
        """
        path = Path(docx_path)
        if not path.exists():
            raise FileNotFoundError(f"Document non trouvé: {path}")

        log.info("gri.parser.start", doc_path=str(path))

        doc = DocxDocument(str(path))
        sections = self._extract_sections(doc)
        tables = self._extract_tables(doc, sections)

        log.info(
            "gri.parser.complete",
            sections=len(sections),
            tables=len(tables),
        )

        return sections, tables

    def _extract_sections(self, doc: Document) -> list[ParsedSection]:
        """Extrait toutes les sections avec leur hiérarchie."""
        sections: list[ParsedSection] = []
        current_section: ParsedSection | None = None
        content_buffer: list[str] = []

        for i, para in enumerate(doc.paragraphs):
            heading_level = self._get_heading_level(para)

            if heading_level:
                # Sauvegarder la section précédente
                if current_section:
                    current_section.content = "\n".join(content_buffer).strip()
                    current_section.end_index = i - 1
                    if current_section.content or current_section.title:
                        sections.append(current_section)

                # Mettre à jour la hiérarchie
                self._update_hierarchy(heading_level, para.text.strip())

                # Détecter le type de section
                section_type = self._detect_section_type(para.text)

                # Créer nouvelle section
                current_section = ParsedSection(
                    level=heading_level,
                    title=para.text.strip(),
                    hierarchy=self._current_hierarchy.copy(),
                    section_type=section_type,
                    start_index=i,
                )
                content_buffer = []

            elif current_section:
                # Ajouter le texte au buffer
                text = para.text.strip()
                if text:
                    content_buffer.append(text)

        # Sauvegarder la dernière section
        if current_section:
            current_section.content = "\n".join(content_buffer).strip()
            current_section.end_index = len(doc.paragraphs) - 1
            sections.append(current_section)

        log.info("gri.parser.sections_extracted", count=len(sections))
        return sections

    def _extract_tables(
        self,
        doc: Document,
        sections: list[ParsedSection],
    ) -> list[ParsedTable]:
        """Extrait toutes les tables du document."""
        tables: list[ParsedTable] = []

        for i, table in enumerate(doc.tables):
            parsed = self._parse_table(table, i)

            # Associer à la section parente
            parent = self._find_parent_section(i, sections, doc)
            if parent:
                parsed.parent_section = parent.title

            # Détecter si c'est une table de jalon
            if self._is_milestone_table(parsed):
                parsed.table_type = "milestone_criteria"
                parsed.milestone_id = extract_milestone_id(parsed.full_text)

            tables.append(parsed)

        log.info("gri.parser.tables_extracted", count=len(tables))
        return tables

    def _parse_table(self, table: Table, index: int) -> ParsedTable:
        """Parse une table DOCX en structure ParsedTable."""
        rows_data: list[dict[str, str]] = []
        headers: list[str] = []
        text_parts: list[str] = []

        for row_idx, row in enumerate(table.rows):
            cells = [self._get_cell_text(cell) for cell in row.cells]

            if row_idx == 0:
                # Première ligne = headers
                headers = [c for c in cells if c]
                text_parts.append(" | ".join(headers))
            else:
                # Lignes de données
                if any(cells):
                    row_dict = {}
                    for col_idx, cell_text in enumerate(cells):
                        if col_idx < len(headers):
                            row_dict[headers[col_idx]] = cell_text
                        else:
                            row_dict[f"col_{col_idx}"] = cell_text
                    rows_data.append(row_dict)
                    text_parts.append(" | ".join(cells))

        return ParsedTable(
            table_index=index,
            headers=headers,
            rows=rows_data,
            full_text="\n".join(text_parts),
        )

    def _get_cell_text(self, cell) -> str:
        """Extrait le texte d'une cellule (gère les cellules fusionnées)."""
        return " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())

    def _get_heading_level(self, para: Paragraph) -> int | None:
        """Détermine le niveau de heading (1-7) ou None si pas un heading."""
        style_name = para.style.name if para.style else ""

        # Heading 1, Heading 2, etc.
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading ", "").strip())
                return min(level, 7)
            except ValueError:
                pass

        # Titre, Titre 1, etc. (français)
        if "Titre" in style_name:
            try:
                parts = style_name.split()
                if len(parts) > 1:
                    return min(int(parts[1]), 7)
                return 1
            except (ValueError, IndexError):
                return 1

        # Détecter via le formatage XML
        outline_level = self._get_outline_level(para)
        if outline_level is not None:
            return outline_level

        return None

    def _get_outline_level(self, para: Paragraph) -> int | None:
        """Récupère le niveau outline depuis le XML."""
        try:
            pPr = para._element.find(qn("w:pPr"))
            if pPr is not None:
                outlineLvl = pPr.find(qn("w:outlineLvl"))
                if outlineLvl is not None:
                    val = outlineLvl.get(qn("w:val"))
                    if val is not None:
                        return int(val) + 1  # 0-based → 1-based
        except Exception:
            pass
        return None

    def _update_hierarchy(self, level: int, title: str) -> None:
        """Met à jour la hiérarchie courante."""
        # Enregistrer ce niveau
        self._current_levels[level] = title

        # Supprimer les niveaux inférieurs
        levels_to_remove = [l for l in self._current_levels if l > level]
        for l in levels_to_remove:
            del self._current_levels[l]

        # Reconstruire la hiérarchie
        self._current_hierarchy = ["GRI"]
        for l in sorted(self._current_levels.keys()):
            self._current_hierarchy.append(self._current_levels[l])

    def _detect_section_type(self, text: str) -> SectionType:
        """Détecte le type de section depuis le titre."""
        for section_type, pattern in self.SECTION_TYPE_PATTERNS.items():
            if pattern.search(text):
                return SectionType(section_type)
        return SectionType.CONTENT

    def _is_milestone_table(self, table: ParsedTable) -> bool:
        """Vérifie si une table contient des critères de jalon."""
        markers = [
            "Critères",
            "passage",
            "Jalon",
            "CDR",
            "TRR",
            "SRR",
            "PDR",
            "IRR",
            "SAR",
            "ORR",
            "PRR",
            "MNR",
            "ASR",
        ]
        text = table.full_text.lower()
        return any(m.lower() in text for m in markers)

    def _find_parent_section(
        self,
        table_index: int,
        sections: list[ParsedSection],
        doc: Document,
    ) -> ParsedSection | None:
        """Trouve la section parente d'une table."""
        # Approximation: on utilise l'index de la table
        # Une meilleure approche serait de tracker la position réelle
        for section in reversed(sections):
            if section.start_index < table_index * 10:  # Heuristique
                return section
        return sections[0] if sections else None

    def extract_phase_num(self, text: str) -> int | None:
        """Extrait le numéro de phase du texte."""
        match = self.PHASE_NUM_PATTERN.search(text)
        if match:
            num = int(match.group(1))
            return num if 1 <= num <= 7 else None
        return None

    def extract_principle_num(self, text: str) -> int | None:
        """Extrait le numéro de principe du texte."""
        match = self.PRINCIPLE_NUM_PATTERN.search(text)
        if match:
            num = int(match.group(1))
            return num if 1 <= num <= 11 else None
        return None
