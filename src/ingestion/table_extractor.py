"""Extracteur de tables spécialisé pour le GRI.

Les tables de critères de jalons sont la partie la plus critique.
Ne jamais perdre ces données.
"""

from pathlib import Path

from docx import Document as DocxDocument
from docx.table import Table

from src.core.logging import get_logger
from src.core.milestone_utils import extract_milestone_id
from src.ingestion.models import ParsedTable

log = get_logger(__name__)


class GRITableExtractor:
    """Extracteur de tables spécialisé pour le document GRI.

    Détecte automatiquement les tables de critères de jalons
    et les tables de livrables.
    """

    # Marqueurs pour identifier les tables de jalons
    JALON_TABLE_MARKERS = [
        "Critères du passage",
        "Critères de passage",
        "Jalon M",
        "Jalon J",
        "CDR",
        "TRR",
        "SRR",
        "PDR",
        "IRR",
        "SAR",
        "ORR",
        "PRR",
        "MNR",
        "ASR",
        "SFR",
    ]

    # Marqueurs pour les tables de livrables
    DELIVERABLES_MARKERS = [
        "Livrable",
        "Produit",
        "Artefact",
        "Document",
        "Output",
    ]

    # Noms des jalons GRI (aligné avec milestone_retriever.py)
    MILESTONE_NAMES = {
        "M0": "MNS - Mission Needs Statement",
        "M1": "SRR - System Requirements Review",
        "M2": "PDR - Preliminary Design Review",
        "M3": "CDR - Critical Design Review",
        "M4": "IRR - Integration Readiness Review",
        "M5": "TRR - Test Readiness Review",
        "M6": "SAR - System Acceptance Review",
        "M7": "ORR - Operational Readiness Review",
        "M8": "MNR - Mission Needs Review",
        "M9": "System Retirement Review",
        "J1": "Jalon J1 (CIR)",
        "J2": "Jalon J2 (CIR)",
        "J3": "Jalon J3 (CIR)",
        "J4": "Jalon J4 (CIR)",
        "J5": "Jalon J5 (CIR)",
        "J6": "Jalon J6 (CIR)",
    }

    def extract(self, docx_path: str | Path) -> list[ParsedTable]:
        """Extrait toutes les tables du document.

        Args:
            docx_path: Chemin vers le fichier DOCX

        Returns:
            Liste de ParsedTable
        """
        path = Path(docx_path)
        if not path.exists():
            raise FileNotFoundError(f"Document non trouvé: {path}")

        log.info("gri.table_extractor.start", doc_path=str(path))

        doc = DocxDocument(str(path))
        tables: list[ParsedTable] = []

        for i, table in enumerate(doc.tables):
            parsed = self._extract_table(table, i)
            tables.append(parsed)

        # Stats
        milestone_tables = sum(1 for t in tables if t.table_type == "milestone_criteria")
        log.info(
            "gri.table_extractor.complete",
            total_tables=len(tables),
            milestone_tables=milestone_tables,
        )

        return tables

    def _extract_table(self, table: Table, index: int) -> ParsedTable:
        """Extrait une table en structure ParsedTable."""
        headers = self._extract_headers(table)
        rows = self._extract_rows(table, headers)
        full_text = self._table_to_text(table)

        # Déterminer le type de table (passer les headers pour une détection plus précise)
        table_type = self._detect_table_type(full_text, headers)

        # Extraire l'ID du jalon si applicable
        milestone_id = None
        if table_type == "milestone_criteria":
            # Priorité: headers puis full_text
            header_text = " ".join(headers) if headers else ""
            milestone_id = extract_milestone_id(header_text) or extract_milestone_id(full_text)

        return ParsedTable(
            table_index=index,
            table_type=table_type,
            headers=headers,
            rows=rows,
            full_text=full_text,
            milestone_id=milestone_id,
        )

    def _extract_headers(self, table: Table) -> list[str]:
        """Extrait les en-têtes de la table.

        Utilise la position géométrique des cellules (cell._tc) pour gérer
        correctement les cellules fusionnées, évitant la désynchronisation
        entre en-têtes et données.
        """
        if not table.rows:
            return []

        headers = []
        seen_positions = set()

        for cell in table.rows[0].cells:
            # Utiliser la position géométrique pour éviter les doublons de cellules fusionnées
            cell_key = (cell._tc.left, cell._tc.top)
            if cell_key not in seen_positions:
                text = self._clean_cell_text(cell)
                # Toujours ajouter, même si vide (pour garder l'alignement)
                headers.append(text if text else f"col_{len(headers)}")
                seen_positions.add(cell_key)

        return headers

    def _extract_rows(self, table: Table, headers: list[str]) -> list[dict[str, str]]:
        """Extrait les lignes de données."""
        rows: list[dict[str, str]] = []

        for row_idx, row in enumerate(table.rows):
            if row_idx == 0:
                continue  # Skip headers

            cells = []
            seen_texts = set()

            for cell in row.cells:
                text = self._clean_cell_text(cell)
                # Gérer les cellules fusionnées
                cell_key = (cell._tc.left, cell._tc.top)
                if cell_key not in seen_texts:
                    cells.append(text)
                    seen_texts.add(cell_key)

            if any(cells):
                row_dict = {}
                for col_idx, cell_text in enumerate(cells):
                    if col_idx < len(headers):
                        row_dict[headers[col_idx]] = cell_text
                    else:
                        row_dict[f"col_{col_idx}"] = cell_text
                rows.append(row_dict)

        return rows

    def _clean_cell_text(self, cell) -> str:
        """Nettoie le texte d'une cellule."""
        paragraphs = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
        return " ".join(paragraphs)

    def _table_to_text(self, table: Table) -> str:
        """Convertit une table en texte formaté."""
        lines: list[str] = []

        for row in table.rows:
            cells = []
            seen = set()

            for cell in row.cells:
                text = self._clean_cell_text(cell)
                cell_key = (cell._tc.left, cell._tc.top)
                if cell_key not in seen:
                    cells.append(text)
                    seen.add(cell_key)

            if any(cells):
                lines.append(" | ".join(cells))

        return "\n".join(lines)

    def _detect_table_type(
        self,
        full_text: str,
        headers: list[str] | None = None,
    ) -> str:
        """Détecte le type de table en priorisant les en-têtes.

        La détection suit un ordre de priorité pour éviter les faux positifs:
        1. Chercher dans les en-têtes (forte confiance)
        2. Chercher dans la première ligne du texte (titre probable)
        3. Fallback sur le texte complet avec critères stricts
        """
        # Priorité 1: Chercher dans les en-têtes
        if headers:
            header_text = " ".join(headers).lower()
            for marker in self.JALON_TABLE_MARKERS:
                if marker.lower() in header_text:
                    return "milestone_criteria"

        # Priorité 2: Chercher dans la première ligne du texte (souvent le titre)
        first_line = full_text.split("\n")[0].lower() if full_text else ""
        for marker in self.JALON_TABLE_MARKERS:
            if marker.lower() in first_line:
                return "milestone_criteria"

        # Priorité 3: Fallback sur le texte complet (avec critères stricts)
        text_lower = full_text.lower()
        # Requérir au moins 2 marqueurs ou un pattern fort pour éviter les faux positifs
        if "critères" in text_lower and any(
            m.lower() in text_lower for m in ["passage", "jalon"]
        ):
            return "milestone_criteria"

        # Tables de livrables
        for marker in self.DELIVERABLES_MARKERS:
            if marker.lower() in text_lower:
                return "deliverables"

        return "general"

    def get_milestone_table(
        self,
        tables: list[ParsedTable],
        milestone_id: str,
    ) -> ParsedTable | None:
        """Récupère la table d'un jalon spécifique."""
        for table in tables:
            if table.milestone_id == milestone_id:
                return table
        return None

    def format_milestone_criteria(self, table: ParsedTable) -> str:
        """Formate les critères d'un jalon pour l'indexation.

        RÈGLE CRITIQUE: Ne jamais fragmenter les critères d'un jalon.
        Tout doit être dans un seul chunk.
        """
        if not table.milestone_id:
            return table.full_text

        milestone_name = self.MILESTONE_NAMES.get(
            table.milestone_id,
            f"Jalon {table.milestone_id}",
        )

        lines = [
            f"Critères de passage du jalon {table.milestone_id} — {milestone_name}",
            "",
        ]

        # Formater chaque critère
        for i, row in enumerate(table.rows, 1):
            criterion = self._find_criterion_text(row)
            if criterion:
                lines.append(f"{i}. {criterion}")

        return "\n".join(lines)

    def _find_criterion_text(self, row: dict[str, str]) -> str:
        """Trouve le texte du critère dans une ligne avec recherche flexible.

        Stratégie:
        1. Recherche partielle de clé (critère, description, besoin, exigence)
        2. Fallback: prendre la valeur la plus longue (le critère est généralement
           le texte le plus long, contrairement aux IDs qui sont courts)
        """
        # Liste des termes indicateurs d'une colonne de critère
        criterion_terms = ["critère", "criteres", "description", "besoin", "exigence"]

        # 1. Recherche partielle de clé
        for key, value in row.items():
            key_lower = key.lower()
            if any(term in key_lower for term in criterion_terms) and value and len(value) > 5:
                return value

        # 2. Fallback: prendre la valeur la plus longue
        # Le critère est généralement le texte le plus long de la ligne
        if row:
            longest = max(row.values(), key=len, default="")
            if len(longest) > 10:  # Seuil minimum pour éviter les IDs
                return longest

        return ""

    def extract_all_milestones(
        self,
        tables: list[ParsedTable],
    ) -> dict[str, ParsedTable]:
        """Extrait toutes les tables de jalons indexées par ID."""
        milestones: dict[str, ParsedTable] = {}

        for table in tables:
            if table.table_type == "milestone_criteria" and table.milestone_id:
                milestones[table.milestone_id] = table

        log.info(
            "gri.table_extractor.milestones",
            found=list(milestones.keys()),
        )

        return milestones
